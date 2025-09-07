# --- Agno imports ---
from agno.agent import Agent
from agno.tools.file import FileTools
from agno.tools.exa import ExaTools
from dotenv import load_dotenv
from agno.agent import Agent
from agno.tools.reasoning import ReasoningTools
from agno.tools.scrapegraph import ScrapeGraphTools
from agno.models.mistral import MistralChat
import requests

# --- File output ---
from docx import Document
import re
import os


load_dotenv()
mistral_api_key = os.getenv("MISTRAL_API_KEY")

Data_Scraper = Agent(
    name="EU regulation crawler agent",
    # model=OpenAIChat(id=id_openai, api_key=api_key_openai),
    model=MistralChat(
            id="mistral-large-latest",
            api_key=mistral_api_key,
    ),
    tools=[
        ReasoningTools(add_instructions=True),
        FileTools(),
        ScrapeGraphTools(),
        ExaTools(include_domains={"eur-lex.europa.eu", "ec.europa.eu"})
    ],
    show_tool_calls=True,

    description="""
    Specialized agent for crawling and extracting compliance-relevant information about European regulations affecting a specific company and
    industry. Focuses on trusted EU sources such as EUR-Lex and the European Commission website.
    """,

    instructions="""
    You are a compliance-focused data scraping agent. Your role is to gather detailed, structured, and relevant information about a given
    company's regulatory environment in Europe, based on its industry. You are primarily responsible for collecting recent European regulations,
    directives, press releases, and legal documents that may affect the company.

    Follow these steps:

    1. Understand the target company and its industry context from the user input.

    2. Use `ExaTools` to search for recent legal documents, directives, or press releases on:
    - https://eur-lex.europa.eu/homepage.html (for legislation and directives)
    - https://ec.europa.eu/commission/presscorner/home/en (for press releases and regulatory announcements)

    3. Use `scrapeGraphTools` to extract detailed data from any relevant pages found.

    4. Focus on identifying potential compliance areas such as:
    - Data privacy (e.g. GDPR)
    - AI regulation (e.g. EU AI Act)
    - Environmental directives
    - Industry-specific regulations

    5. If no regulations are found that specifically name the company, fall back to summarizing recent EU laws or directives that are
    likely to affect the industry. For example:
    - If no AI laws mention the company, summarize how the EU AI Act impacts the industry generally.

    6. Structure the findings into a clear and detailed Markdown file using `FileTools`. The file should include:
    - A metadata section at the top with:
      - Company name
      - Industry
      - Report generation date
      - List of data sources used
    - A table of contents or logical section headers
    - Categorized summaries (e.g. AI, Sustainability, Data Privacy)
    - Direct source links to each regulation
    - Dates of publication or enforcement

    7. Save the final file using `FileTools` with the name `compliance_[industry].md` (e.g. `compliance_automotive.md`).

    Only return verified and relevant information from the specified EU sources. Ignore speculative or opinion-based content. Prioritize
    legal, and up-to-date documents from official EU regulatory domains.
    """,
    markdown=True,
    stream=True,
    success_criteria="""
    The agent must generate the final file `compliance_automotive.md` using `FileTools`.
    """,
)


# Execute the agent with a specific query
Data_Scraper.print_response(
    "Find the most recent EU directives or regulations that apply to Renault in the automotive sector, "
    "especially related to AI, sustainability, or emissions. Structure the results in a clear compliance "
    "report, and save the final output as `compliance_automotive.md` file",
    markdown=True
)





# Load markdown content
with open(r"C:\Users\user\Desktop\ComplianceAgent\compliance_automotive.md", "r", encoding="utf-8") as file:
    final_output = file.read()

# Initialize Word doc
from docx import Document
import re
import os

doc = Document()
doc.add_heading("Final Compliance Analysis", 0)  # Level 0 for main heading

def add_formatted_paragraph(doc, text):
    """Helper function to add a paragraph with properly formatted bold text"""
    p = doc.add_paragraph()
    # Split text by bold markers (**bold text**)
    parts = re.split(r'(\*\*.+?\*\*)', text)  # Non-greedy match between ** markers
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])  # Remove the ** markers
            run.bold = True
        else:
            p.add_run(part)
    return p

# Process the markdown line by line
for line in final_output.splitlines():
    line = line.strip()
    if not line:
        continue  # Skip empty lines
    
    if line.startswith("# "):
        doc.add_heading(line[2:], level=1)  # H1
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=2)  # H2
    elif line.startswith("### "):
        doc.add_heading(line[4:], level=3)  # H3
    else:
        add_formatted_paragraph(doc, line)

# Save the Word file
output_path = r"C:\Users\user\Desktop\ComplianceAgent\compliance_automotive.docx"
doc.save(output_path)

print(f"✅ DOCX file saved at: {output_path}")
